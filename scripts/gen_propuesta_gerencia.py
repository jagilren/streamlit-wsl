"""Genera un .docx con la propuesta de valor del sistema PTAR para gerencia.

Salida: <root>/PTAR_Propuesta_Gerencia.docx
Ejecutar:  .venv/bin/python scripts/gen_propuesta_gerencia.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

AZUL = RGBColor(0x18, 0x5F, 0xA5)
GRIS = RGBColor(0x55, 0x55, 0x55)
NEGRO = RGBColor(0x00, 0x00, 0x00)

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "PTAR_Propuesta_Gerencia.docx"


def add_titulo(doc: Document, texto: str, size: int = 22) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = AZUL


def add_subtitulo(doc: Document, texto: str, size: int = 12) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.italic = True
    run.font.size = Pt(size)
    run.font.color.rgb = GRIS


def add_seccion(doc: Document, numero: int, titulo: str) -> None:
    p = doc.add_paragraph()
    run_num = p.add_run(f"{numero}. ")
    run_num.bold = True
    run_num.font.size = Pt(14)
    run_num.font.color.rgb = AZUL
    run_tit = p.add_run(titulo)
    run_tit.bold = True
    run_tit.font.size = Pt(14)
    run_tit.font.color.rgb = AZUL
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_parrafo(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(11)
    run.font.color.rgb = NEGRO
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25


def add_bullet(doc: Document, texto: str, negrita_intro: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if negrita_intro:
        run_b = p.add_run(negrita_intro)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_resto = p.add_run(" " + texto)
        run_resto.font.size = Pt(11)
    else:
        run = p.add_run(texto)
        run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)


def main() -> None:
    doc = Document()

    # Márgenes razonables
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── Encabezado ───────────────────────────────────────────────────────────
    add_titulo(doc, "Sistema de Monitoreo PTAR Industrial", size=22)
    add_subtitulo(
        doc,
        "Propuesta de valor para la Gerencia Ambiental y de Planta",
        size=12,
    )
    add_subtitulo(doc, "Red Proyectos con Ingeniería (RPCI) · 2026", size=10)

    doc.add_paragraph()

    add_parrafo(
        doc,
        "El siguiente documento resume las seis razones por las que implementar "
        "el sistema de monitoreo en línea de la PTAR (DQO, pH, Oxígeno Disuelto "
        "y operaciones complementarias) genera retorno medible en cumplimiento "
        "normativo, operación y reputación corporativa.",
    )

    # ── 1. Cumplimiento normativo y reducción de riesgo legal ────────────────
    add_seccion(doc, 1, "Cumplimiento normativo continuo y trazabilidad legal")
    add_parrafo(
        doc,
        "El sistema registra cada lectura de los parámetros críticos cada 3 "
        "minutos, con sello de tiempo y almacenamiento permanente en una base "
        "de datos histórica (TimescaleDB). Esto convierte la planta en una "
        "operación auditable en cualquier momento.",
    )
    add_bullet(
        doc,
        "Evidencia objetiva ante CAR, ANLA y auditorías ISO 14001 — el sistema "
        "construye automáticamente el registro de cumplimiento exigido por la "
        "resolución de vertimientos aplicable.",
        negrita_intro="Defensa ante autoridades:",
    )
    add_bullet(
        doc,
        "Las desviaciones de DQO, pH u OD quedan registradas con timestamp y "
        "magnitud, eliminando la zona gris entre 'no lo vimos' y 'no pasó'.",
        negrita_intro="Trazabilidad sin lagunas:",
    )
    add_bullet(
        doc,
        "Sanciones por superar límites de vertimiento pueden alcanzar miles de "
        "SMMLV. Un solo evento evitado paga el sistema varias veces.",
        negrita_intro="Reducción de multas:",
    )

    # ── 2. Detección temprana de eventos ──────────────────────────────────────
    add_seccion(doc, 2, "Detección temprana de eventos fuera de rango")
    add_parrafo(
        doc,
        "Cada dashboard clasifica las lecturas en tiempo real en óptimo, "
        "advertencia y crítico, con indicadores visuales pulsantes ('EN VIVO') "
        "y tablas de eventos. El operador ve la desviación en el momento en "
        "que ocurre, no al día siguiente cuando ya es irreversible.",
    )
    add_bullet(
        doc,
        "Si el pH se acerca al umbral crítico o el OD del reactor biológico "
        "cae por debajo del mínimo, la condición se visualiza inmediatamente "
        "y permite intervenir antes de que la biomasa colapse.",
        negrita_intro="Ventana de reacción:",
    )
    add_bullet(
        doc,
        "Cada panel se actualiza automáticamente cada 60 segundos sin "
        "intervención humana; no depende de rondas manuales ni de Excel.",
        negrita_intro="Sin dependencia de turnos:",
    )

    # ── 3. Decisiones basadas en datos, no en intuición ──────────────────────
    add_seccion(doc, 3, "Decisiones operativas basadas en datos verificables")
    add_parrafo(
        doc,
        "Los KPIs ejecutivos (DQO afluente, DQO efluente, eficiencia de "
        "remoción y cumplimiento normativo) se calculan automáticamente sobre "
        "los rangos seleccionados — hoy, esta semana, este mes, últimos 90 "
        "días o el año en curso — con comparación contra el período anterior.",
    )
    add_bullet(
        doc,
        "La gerencia puede medir si una modificación operativa (dosificación, "
        "tiempos de aireación, recirculación) realmente mejoró la eficiencia "
        "de remoción o solo lo aparentó.",
        negrita_intro="Antes / después medible:",
    )
    add_bullet(
        doc,
        "Las gráficas de tendencia con bandas óptimas sombreadas permiten ver "
        "patrones (oscilaciones diarias, derivas, eventos repetitivos) que "
        "en planillas de Excel son invisibles.",
        negrita_intro="Visualización de patrones:",
    )

    # ── 4. Optimización de costos operativos ─────────────────────────────────
    add_seccion(doc, 4, "Optimización de costos químicos y energéticos")
    add_parrafo(
        doc,
        "Conocer en línea el comportamiento real del proceso permite ajustar "
        "consumos en lugar de sobre-dosificar 'por seguridad'.",
    )
    add_bullet(
        doc,
        "Si el OD se mantiene sostenidamente por encima del óptimo, hay "
        "potencial de reducir caudal de aire — los sopladores son típicamente "
        "el mayor consumidor eléctrico de la PTAR.",
        negrita_intro="Aireación ajustada:",
    )
    add_bullet(
        doc,
        "Dosificación de químicos correctivos de pH guiada por la tendencia "
        "real, no por consignas conservadoras.",
        negrita_intro="Reactivos químicos:",
    )
    add_bullet(
        doc,
        "Mantenimiento predictivo: las desviaciones sostenidas de un "
        "transmisor sugieren deriva del sensor y permiten programar "
        "calibración antes de que falle.",
        negrita_intro="Mantenimiento dirigido:",
    )

    # ── 5. Centralización y acceso multiusuario ──────────────────────────────
    add_seccion(doc, 5, "Plataforma centralizada, multiusuario y multidispositivo")
    add_parrafo(
        doc,
        "Todos los módulos (DQO, pH, OD y los que se incorporen a futuro) "
        "viven bajo una sola URL con autenticación por usuario. El gerente "
        "ambiental, el jefe de planta y el operador ven la misma información "
        "actualizada — sin reportes pegados por correo a fin de mes.",
    )
    add_bullet(
        doc,
        "Encabezado fijo con el reloj, el módulo activo y el usuario logueado; "
        "se ve igual en computador de escritorio que en celular en planta.",
        negrita_intro="Acceso desde cualquier lugar:",
    )
    add_bullet(
        doc,
        "Filtros de tiempo compartidos: el rango seleccionado en un dashboard "
        "persiste al navegar a otro, permitiendo análisis cruzado entre "
        "parámetros (¿cayó el OD el mismo día que subió la DQO?).",
        negrita_intro="Análisis cruzado:",
    )

    # ── 6. Escalabilidad y futuro ────────────────────────────────────────────
    add_seccion(doc, 6, "Arquitectura escalable y a prueba de obsolescencia")
    add_parrafo(
        doc,
        "El sistema fue diseñado para crecer sin reescribirse. Cada operación "
        "nueva (caudal, sólidos suspendidos, conductividad, nitrógeno, etc.) "
        "se incorpora como un módulo independiente que hereda el patrón ya "
        "validado: autenticación, filtros, autorefresh y persistencia.",
    )
    add_bullet(
        doc,
        "TimescaleDB en contenedor Docker, código en Python con bibliotecas "
        "abiertas (Streamlit, Plotly, psycopg2). Sin licencias atadas a "
        "proveedores propietarios.",
        negrita_intro="Sin lock-in tecnológico:",
    )
    add_bullet(
        doc,
        "Cada módulo es un microservicio que corre de forma autónoma, así "
        "una falla en pH no afecta el monitoreo de DQO ni de OD.",
        negrita_intro="Resiliencia operativa:",
    )
    add_bullet(
        doc,
        "El histórico queda en una base estándar SQL; puede integrarse con "
        "Power BI, ERP o sistemas SCADA si la empresa lo requiere a futuro.",
        negrita_intro="Integración futura:",
    )

    # ── Cierre ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    add_seccion(doc, 0, "")  # spacer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Implementar este sistema no es un gasto: es un seguro contra "
        "sanciones, una herramienta de optimización y una ventaja competitiva "
        "frente a auditorías ambientales y clientes que exigen evidencia "
        "de gestión responsable del agua."
    )
    run.italic = True
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = AZUL

    doc.save(OUT)
    print(f"OK → {OUT}")


if __name__ == "__main__":
    main()
